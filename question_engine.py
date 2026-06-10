"""
question_engine.py — standalone question generation for the K-12 RAG system.

No Streamlit dependency; safe to import in tests, notebooks, or CLI scripts.

Also owns the shared document-loading pipeline used by both app.py and
run_experiment.py:

    load_docs_from_data_dir()   → List[Document]
    build_vectorstore()         → FAISS  (creates + optionally saves index)
"""

import io
import json
import logging
from pathlib import Path
from typing import Any, Dict, List, Optional

from langchain_core.documents import Document
from langchain_ollama import ChatOllama
from langchain_text_splitters import RecursiveCharacterTextSplitter

logger = logging.getLogger(__name__)

DEFAULT_MODEL      = "llama3.2"
DEFAULT_OLLAMA_URL = "http://localhost:11434"
DATA_DIR           = Path(__file__).parent / "data"

_RETRIEVAL_QUERY = "main concepts key ideas overview"
_RETRIEVAL_K     = 8
_CHUNK_PREVIEW   = 600


# ── Document loading (shared with app.py) ─────────────────────────────────────

def _extract_pdf_text(data: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(data))
    pages = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            pages.append("")
    return "\n".join(pages)


def _extract_docx_text(data: bytes) -> str:
    from docx import Document as DocxDocument
    doc = DocxDocument(io.BytesIO(data))
    parts: List[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


def file_to_documents(
    name: str,
    data: bytes,
    chunk_size: int = 1000,
    chunk_overlap: int = 150,
) -> List[Document]:
    """Convert raw file bytes to a list of chunked Documents.

    Handles: .pdf, .txt, .md, .docx
    Skips temp lock files (names starting with ~$).
    """
    if name.startswith("~$"):
        return []

    ext = (name.rsplit(".", 1)[-1] or "").lower()
    if ext == "pdf":
        content = _extract_pdf_text(data)
    elif ext in ("txt", "md"):
        content = data.decode("utf-8", errors="ignore")
    elif ext == "docx":
        content = _extract_docx_text(data)
    else:
        content = data.decode("utf-8", errors="ignore")

    if not content.strip():
        return []

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", " ", ""],
    )
    return [
        Document(page_content=c, metadata={"source": name, "chunk": i})
        for i, c in enumerate(splitter.split_text(content))
    ]


def load_docs_from_data_dir(
    data_dir: Optional[Path] = None,
    chunk_size: int = 1000,
    chunk_overlap: int = 150,
) -> List[Document]:
    """Load and chunk all supported documents from data_dir (default: data/).

    Supported formats: .pdf  .txt  .md  .docx
    Temp Word lock files (~$...) are automatically skipped.
    """
    data_dir = data_dir or DATA_DIR
    docs: List[Document] = []
    for path in sorted(data_dir.iterdir()):
        if path.suffix.lower() in (".pdf", ".txt", ".md", ".docx"):
            docs.extend(file_to_documents(path.name, path.read_bytes(), chunk_size, chunk_overlap))
    return docs


def build_vectorstore(
    docs: List[Document],
    embed_model: str = "nomic-embed-text",
    ollama_url: Optional[str] = None,
    save_path: Optional[Path] = None,
):
    """Embed docs with OllamaEmbeddings and build a FAISS vectorstore.

    If save_path is provided the index is saved to disk so run_experiment.py
    and app.py can reload it without re-embedding.
    """
    from langchain_community.vectorstores import FAISS
    from langchain_ollama import OllamaEmbeddings

    url        = ollama_url or DEFAULT_OLLAMA_URL
    embeddings = OllamaEmbeddings(model=embed_model, base_url=url)
    vs         = FAISS.from_documents(docs, embeddings)
    if save_path is not None:
        vs.save_local(str(save_path))
    return vs

GRADE_DESCRIPTIONS: Dict[str, str] = {
    "K-2":  "kindergarten to grade 2 (ages 5–7). Use very simple words, short sentences, and concrete examples.",
    "3-5":  "grades 3 to 5 (ages 8–10). Use clear language and relatable everyday examples.",
    "6-8":  "grades 6 to 8 (ages 11–13). Use moderate vocabulary and expect some prior knowledge.",
    "9-12": "grades 9 to 12 (ages 14–18). Use academic vocabulary and expect abstract reasoning.",
}

BLOOM_DESCRIPTIONS: Dict[str, str] = {
    "Remember":   "recall facts and basic concepts (e.g., define, list, name, identify)",
    "Understand": "explain ideas in own words (e.g., describe, summarise, explain, classify)",
    "Apply":      "use information in new situations (e.g., solve, demonstrate, use, calculate)",
    "Analyze":    "draw connections and break down information (e.g., compare, contrast, differentiate)",
}

_MCQ_SCHEMA   = '{"question": "...", "options": {"a": "...", "b": "...", "c": "...", "d": "..."}, "correct": "a", "explanation": "..."}'
_SHORT_SCHEMA = '{"question": "...", "sample_answer": "...", "explanation": "..."}'
_TF_SCHEMA    = '{"question": "...", "correct": "True" or "False", "explanation": "..."}'

_RETRY_SUFFIX = (
    "\n\nIMPORTANT: Your previous response could not be parsed as a JSON array. "
    "Return ONLY a raw JSON array — start with [ and end with ]. "
    "No markdown fences, no extra keys, no commentary outside the array."
)


def _build_prompt(
    context: Optional[str],
    n: int,
    q_type: str,
    grade_band: str,
    bloom_level: str,
) -> str:
    """Build the generation prompt.

    When context is None the context block is omitted entirely (zero-shot
    condition). The rest of the prompt — instructions, grade/Bloom's
    descriptions, JSON schema, rules — is identical in both cases.
    """
    grade_desc = GRADE_DESCRIPTIONS[grade_band]
    bloom_desc = BLOOM_DESCRIPTIONS[bloom_level]

    if q_type == "MCQ":
        schema = _MCQ_SCHEMA
        type_instruction = "multiple-choice questions with 4 options (a, b, c, d) and one correct answer"
    elif q_type == "Short Answer":
        schema = _SHORT_SCHEMA
        type_instruction = "short-answer questions with a sample answer (2–4 sentences)"
    else:  # True/False
        schema = _TF_SCHEMA
        type_instruction = "true/false questions"

    if context is not None:
        context_block = (
            "Use ONLY the information in the context below. Do not use external knowledge.\n\n"
            f"Context:\n{context}\n\n"
        )
    else:
        context_block = ""

    return (
        "You are a K-12 curriculum expert creating assessment questions.\n\n"
        f"Your task: Generate exactly {n} {type_instruction}.\n"
        f"Target audience: {grade_desc}\n"
        f"Bloom's Taxonomy level: {bloom_level} — {bloom_desc}\n\n"
        f"{context_block}"
        "Rules:\n"
        "- Questions must be appropriate for the grade band in vocabulary and complexity.\n"
        "- Each question must target the specified Bloom's level.\n"
        "- Do not repeat questions.\n"
        "- Return ONLY a valid JSON array. No markdown, no explanation outside the JSON.\n\n"
        f"JSON schema for each item:\n{schema}\n\n"
        f"Return a JSON array of exactly {n} items."
    )


def _parse_response(raw: str) -> List[dict]:
    """Parse LLM output into a list of question dicts.

    Accepts either a JSON array OR a single JSON object (llama3.2 with
    format='json' consistently returns one object regardless of n= in the
    prompt). A single object is wrapped in a one-element list.

    Strips markdown code fences defensively, then validates structure.

    Raises:
        json.JSONDecodeError: output is not valid JSON.
        ValueError: output is valid JSON but not a dict or list of dicts.
    """
    text = raw.strip()
    if text.startswith("```"):
        parts = text.split("```")
        text = parts[1] if len(parts) > 1 else text
        if text.startswith("json"):
            text = text[4:]
        text = text.strip()

    data = json.loads(text)

    if isinstance(data, dict):
        # Single question object — model returned one item instead of an array.
        return [data]
    if not isinstance(data, list):
        raise ValueError(f"Expected a JSON object or array, got {type(data).__name__}: {text[:120]}")
    if data and not isinstance(data[0], dict):
        raise ValueError(f"Array elements must be objects, got {type(data[0]).__name__}")

    return data


def generate_questions(
    grade_band: str,
    bloom_level: str,
    question_type: str,
    n: int,
    retrieval: bool = True,
    model: Optional[str] = None,
    ollama_url: Optional[str] = None,
    temperature: float = 0.2,
    vs=None,
) -> Dict[str, Any]:
    """Generate K-12 assessment questions, optionally grounded in FAISS retrieval.

    Args:
        grade_band:    One of "K-2", "3-5", "6-8", "9-12".
        bloom_level:   One of "Remember", "Understand", "Apply", "Analyze".
        question_type: One of "MCQ", "Short Answer", "True/False".
        n:             Number of questions to generate.
        retrieval:     True  → RAG condition: FAISS similarity search grounds the prompt.
                       False → zero-shot condition: same prompt, no context block.
        model:         Ollama chat model name. Defaults to DEFAULT_MODEL.
        ollama_url:    Ollama base URL. Defaults to DEFAULT_OLLAMA_URL.
        temperature:   LLM sampling temperature.
        vs:            FAISS vectorstore instance. Required when retrieval=True.

    Returns:
        {
            "questions":        list[dict],  # parsed question objects ([] on total failure)
            "retrieved_chunks": list[dict],  # chunks used as context ([] when retrieval=False)
            "raw_prompt":       str,         # prompt sent on the last attempt
            "parse_failed":     bool,        # True if both parse attempts failed
            "model":            str,
            "retrieval":        bool,
        }

    Raises:
        ValueError: if retrieval=True and vs is None.
    """
    model = model or DEFAULT_MODEL
    url = ollama_url or DEFAULT_OLLAMA_URL

    if retrieval and vs is None:
        raise ValueError("vs (vectorstore) must be provided when retrieval=True")

    # ── Retrieval ────────────────────────────────────────────────────────────
    retrieved_chunks: List[dict] = []
    context: Optional[str] = None

    if retrieval:
        sample_docs: List[Document] = vs.similarity_search(_RETRIEVAL_QUERY, k=_RETRIEVAL_K)
        for d in sample_docs:
            retrieved_chunks.append({
                "source":  d.metadata.get("source", ""),
                "chunk":   d.metadata.get("chunk", ""),
                "content": d.page_content[:_CHUNK_PREVIEW],
            })
        context = "\n\n".join(c["content"] for c in retrieved_chunks)

    # ── Build prompt and LLM ─────────────────────────────────────────────────
    # Ask for one question per call: llama3.2 with format="json" reliably
    # returns a single JSON object and ignores array-of-N instructions.
    # We loop n times and collect individual results.
    single_prompt = _build_prompt(context, 1, question_type, grade_band, bloom_level)
    retry_prompt  = single_prompt + _RETRY_SUFFIX
    llm = ChatOllama(model=model, temperature=temperature, base_url=url, format="json")

    questions: List[dict] = []
    per_q_failures = 0

    for q_idx in range(1, n + 1):
        # ── Attempt 1 ────────────────────────────────────────────────────────
        raw = llm.invoke(single_prompt).content.strip()
        logger.debug("Q%d attempt 1 [model=%s retrieval=%s]: %s",
                     q_idx, model, retrieval, raw[:200])
        try:
            questions.extend(_parse_response(raw))
            logger.debug("Q%d attempt 1 OK", q_idx)
            continue
        except (json.JSONDecodeError, ValueError) as exc:
            logger.warning("Q%d attempt 1 failed [model=%s retrieval=%s]: %s | raw=%s",
                           q_idx, model, retrieval, exc, raw[:200])

        # ── Attempt 2 — retry with clarification ─────────────────────────────
        raw2 = llm.invoke(retry_prompt).content.strip()
        logger.debug("Q%d attempt 2 [model=%s retrieval=%s]: %s",
                     q_idx, model, retrieval, raw2[:200])
        try:
            questions.extend(_parse_response(raw2))
            logger.info("Q%d attempt 2 OK (retry) [model=%s retrieval=%s]",
                        q_idx, model, retrieval)
        except (json.JSONDecodeError, ValueError) as exc:
            logger.error("Q%d attempt 2 failed [model=%s retrieval=%s]: %s | raw=%s",
                         q_idx, model, retrieval, exc, raw2[:200])
            per_q_failures += 1

    parse_failed = len(questions) == 0
    if questions:
        logger.info("Cell complete [model=%s retrieval=%s]: %d/%d questions, %d failures",
                    model, retrieval, len(questions), n, per_q_failures)

    return {
        "questions":        questions,
        "retrieved_chunks": retrieved_chunks,
        "raw_prompt":       single_prompt,
        "parse_failed":     parse_failed,
        "model":            model,
        "retrieval":        retrieval,
    }
